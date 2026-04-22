import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open("evidence_data.json", encoding="utf-8") as f:
    records = json.load(f)

doc = Document()

title = doc.add_heading("AI Intake Triage — Evidence Log", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Estate Planning Client Triage System — Groq / llama-3.3-70b-versatile")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

for i, record in enumerate(records, 1):
    client_name = record.get("input_data", {}).get("name", f"Client {i}")
    doc.add_heading(f"Client {i}: {client_name}", level=1)

    # 1. Input Data
    doc.add_heading("1. Input Data", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for run in hdr[0].paragraphs[0].runs + hdr[1].paragraphs[0].runs:
        run.bold = True

    for field, value in record["input_data"].items():
        row = table.add_row()
        row.cells[0].text = field
        row.cells[1].text = str(value)

    doc.add_paragraph()

    # 2. Exact Prompt Sent to LLM
    doc.add_heading("2. Exact Prompt Sent to LLM", level=2)

    p = doc.add_paragraph()
    p.add_run("SYSTEM PROMPT:").bold = True
    doc.add_paragraph(record["prompt_sent"]["system"])

    p2 = doc.add_paragraph()
    p2.add_run("USER PROMPT:").bold = True
    doc.add_paragraph(record["prompt_sent"]["user"])

    # 3. Raw LLM Response (Unmodified)
    doc.add_heading("3. Raw LLM Response (Unmodified)", level=2)
    doc.add_paragraph(record["raw_llm_response"])

    # 4. Final Parsed Output
    doc.add_heading("4. Final Parsed Output (JSON)", level=2)
    doc.add_paragraph(json.dumps(record["parsed_output"], indent=2))

    doc.add_page_break()

doc.save("evidence_log.docx")
print(f"evidence_log.docx generated — {len(records)} client(s) documented.")
